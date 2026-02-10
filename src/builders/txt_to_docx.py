from __future__ import annotations

from pathlib import Path
import re
from docx import Document

TXT_PATH = Path("out.txt")
DOCX_PATH = Path("contract.docx")

# "1. ОПРЕДЕЛЕНИЯ"
_RE_SECTION_HEADING = re.compile(r"^\s*(\d+)\.\s+\S.+$")

# "1.1. текст" / "10.12. текст"
_RE_SUBCLAUSE = re.compile(r"^\s*\d+\.\d+(?:\.\d+)*\.\s+\S.+$")

# "ARTICLE 1" / "SECTION 1"
_RE_ARTICLE_SECTION = re.compile(r"^\s*(ARTICLE|SECTION)\b", re.IGNORECASE)


def txt_to_docx(txt_path: Path, docx_path: Path) -> None:
    text = txt_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    # --- helper: safe style apply ---
    def add_paragraph(line: str, style: str | None = None):
        p = doc.add_paragraph(line)
        if style:
            try:
                p.style = style
            except Exception:
                # если стиля нет в Word-шаблоне — оставим как есть
                pass
        return p

    # 1) find first non-empty line (title)
    first_nonempty_idx = None
    for i, ln in enumerate(lines):
        if ln.strip():
            first_nonempty_idx = i
            break

    for i, raw in enumerate(lines):
        line = (raw or "").strip()
        if not line:
            continue

        # Title (first non-empty)
        if first_nonempty_idx is not None and i == first_nonempty_idx:
            if line.upper().startswith(("ДОГОВОР", "КОНТРАКТ", "CONTRACT", "AGREEMENT")):
                add_paragraph(line, "Title")
                continue

        # ARTICLE / SECTION headings (EN)
        if _RE_ARTICLE_SECTION.match(line):
            add_paragraph(line, "Heading 1")
            continue

        # Section heading "1. ..."
        # IMPORTANT: do not treat "1.1. ..." as heading
        if _RE_SECTION_HEADING.match(line) and not _RE_SUBCLAUSE.match(line):
            add_paragraph(line, "Heading 1")
            continue

        # Subclause "1.1. ..."
        if _RE_SUBCLAUSE.match(line):
            add_paragraph(line, "Normal")
            continue

        # Default
        add_paragraph(line, "Normal")

    doc.save(docx_path)


if __name__ == "__main__":
    txt_to_docx(TXT_PATH, DOCX_PATH)
    print(f"DOCX saved: {DOCX_PATH}")
