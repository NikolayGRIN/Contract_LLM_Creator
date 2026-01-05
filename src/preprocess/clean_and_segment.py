from __future__ import annotations

from docx import Document
from pathlib import Path
import json
import re
import argparse


# ---------------------------
# 1) SETTINGS
# ---------------------------
# Ограничиваем глубину нумерации заголовков:
# MAX_LEVEL=1  -> 1., 2., 3.
# MAX_LEVEL=2  -> 1.1, 2.3, 10.2
MAX_LEVEL = 2

# Заголовок без нумерации (ALL CAPS) считаем заголовком,
# если он достаточно короткий и похож на "GENERAL TERMS..."
ALL_CAPS_MAX_LEN = 90
ALL_CAPS_MIN_ALPHA = 6  # минимум букв, чтобы не ловить мусор


# ---------------------------
# 2) CLEANING
# ---------------------------
def clean_text(text: str) -> str:
    """
    Remove obvious noise but KEEP line structure.
    Segmentation relies on newlines.
    """
    text = re.sub(r"Page\s+\d+\s+of\s+\d+", "", text, flags=re.IGNORECASE)

    text = "\n".join(
        re.sub(r"\s{2,}", " ", line).strip()
        for line in text.splitlines()
    )

    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


# ---------------------------
# 3) HEADING DETECTION
# ---------------------------

PLAIN_HEADINGS = {
    # RU
    "ПРЕАМБУЛА", "Преамбула",
    "ПРЕДМЕТ КОНТРАКТА", "Предмет Контракта",
    "ОСНОВНЫЕ УСЛОВИЯ", "Основные условия",
    "УСЛОВИЯ ОПЛАТЫ", "Условия оплаты",
    "ОТГРУЗКА ТОВАРА", "Отгрузка Товара",
    "ОТВЕТСТВЕННОСТЬ СТОРОН", "Ответственность Сторон",
    "АРБИТРАЖ", "Арбитраж",
    "ПРОЧИЕ УСЛОВИЯ", "Прочие условия",
    "РЕКВИЗИТЫ", "Реквизиты",
    "ПОДПИСИ", "Подписи",
    "ПРИЛОЖЕНИЯ", "Приложения", "Приложения:",

    # EN
    "PREAMBLE", "Preamble",
    "SUBJECT OF THE CONTRACT", "Subject of the Contract",
    "PRINCIPAL PROVISIONS", "Principal provisions",
    "TERMS AND PROCEDURES OF PAYMENT", "Terms and procedures of payment",
    "SHIPMENT OF THE GOODS", "Shipment of the Goods",
    "LIABILITY OF THE PARTIES", "Liability of the Parties",
    "ARBITRATION", "Arbitration",
    "OTHER CONDITIONS", "Other conditions",
    "SIGNATURES", "Signatures",
    "APPENDICES", "Appendices", "Appendices:",
}

ARTICLE_RE = re.compile(r"^(ARTICLE|CLAUSE|SECTION)\s+\d+(\.\d+)*[\.\)]?(\s+.+)?$", re.IGNORECASE)
ADDENDUM_RE = re.compile(r"^(Приложение|Addendum)\s*№?\s*\d+.*$", re.IGNORECASE)
NUMERIC_RE = re.compile(r"^\d+(\.\d+)*[.)]\s+.+$")

TITLE_KEYWORDS = {
    # RU
    "замечания", "качество", "количество", "ассортимент", "комплектн",
    "предмет", "оплата", "поставка", "отгрузка", "ответственность", "арбитраж",
    "конфиденциал", "форс", "обстоятельств",

    # EN
    "subject", "payment", "delivery", "shipment", "liability", "arbitration",
    "quantity", "quality", "selection", "completeness",
    "force majeure", "governing law", "confidential",
    "terms and conditions", "general terms",
}

ACTION_VERBS = {
    # RU
    "принять", "предоставить", "оплатить", "возместить", "осуществить", "направить", "произвести", "передать",
    "просим", "требуем", "заменить",

    # EN
    "accept", "provide", "pay", "return", "compensate", "perform", "send",
    "deliver", "please", "replace",
}


def normalize_spaces(s: str) -> str:
    return re.sub(r"\s{2,}", " ", (s or "").strip())


def numbering_level(line: str) -> int:
    """
    '1. ...' -> 1
    '4.1. ...' -> 2
    '3.1.1. ...' -> 3
    """
    m = re.match(r"^(\d+(?:\.\d+)*)[.)]\s+", line.strip())
    if not m:
        return 0
    return len(m.group(1).split("."))


def starts_with_action_verb(line: str) -> bool:
    s = (line or "").strip()
    if not s:
        return False
    s2 = re.sub(r"^\d+(\.\d+)*[.)]\s+", "", s).strip().lower()
    first = s2.split(" ", 1)[0] if s2 else ""
    return first in ACTION_VERBS


def looks_like_real_title(line: str) -> bool:
    """
    Heuristic: real titles often end with ':' OR contain title keywords.
    """
    s = (line or "").strip()
    if not s:
        return False
    if s.endswith(":"):
        return True
    low = s.lower()
    if any(k in low for k in TITLE_KEYWORDS):
        return True
    return False

NON_SECTION_CAPS = {
    "BUYER", "SELLER", "ПОКУПАТЕЛЬ", "ПОСТАВЩИК", "ПРОДАВЕЦ",
    "CONTRACT", "КОНТРАКТ",
    "CONTRACT №", "КОНТРАКТ №",
    "EXAMPLE", "EXAMPLE:",
}


def is_all_caps_heading(line: str) -> bool:
    s = normalize_spaces(line)
    if not s:
        return False

    # исключаем титульные поля
    for bad in NON_SECTION_CAPS:
        if s.startswith(bad):
            return False

    # не считаем строки в скобках заголовками
    if s.startswith("(") and s.endswith(")"):
        return False

    if len(s) > ALL_CAPS_MAX_LEN:
        return False

    letters = re.findall(r"[A-Za-zА-ЯЁа-яё]", s)
    if len(letters) < ALL_CAPS_MIN_ALPHA:
        return False

    return all(ch.isupper() for ch in letters)



def is_heading(line: str, *, allow_all_caps: bool = True) -> bool:
    s = normalize_spaces(line)
    if not s:
        return False

    if s in PLAIN_HEADINGS:
        return True

    if ARTICLE_RE.match(s):
        return True

    if ADDENDUM_RE.match(s):
        return True

    # Numeric headings with depth control
    if NUMERIC_RE.match(s):
        lvl = numbering_level(s)
        if lvl == 0 or lvl > MAX_LEVEL:
            return False

        # list-actions should NOT become sections
        if starts_with_action_verb(s):
            return False

        return looks_like_real_title(s)

    # All caps headings (for documents without numbering)
    if allow_all_caps and is_all_caps_heading(s):
        return True

    return False


# ---------------------------
# 4) SPLIT INTO SECTIONS
# ---------------------------
def split_into_sections(text: str):
    lines = [normalize_spaces(ln) for ln in text.splitlines() if ln.strip()]

    sections = []
    current_title = "PREFACE"
    current_body: list[str] = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # Combine RU + EN headings if they go one after another
        if (
            (line in PLAIN_HEADINGS or bool(ADDENDUM_RE.match(line)) or is_all_caps_heading(line))
            and i + 1 < len(lines)
            and is_heading(lines[i + 1], allow_all_caps=True)
            and line.lower() != lines[i + 1].lower()
        ):
            combined = f"{line} / {lines[i + 1]}"
            if current_body:
                sections.append({"section": current_title, "text": "\n".join(current_body).strip()})
            current_title = combined
            current_body = []
            i += 2
            continue

        # Normal heading start
        if is_heading(line, allow_all_caps=(current_title == "PREFACE")):
            if current_body:
                sections.append({"section": current_title, "text": "\n".join(current_body).strip()})
            current_title = line
            current_body = []
        else:
            current_body.append(line)

        i += 1

    if current_body:
        sections.append({"section": current_title, "text": "\n".join(current_body).strip()})

    if not sections and text.strip():
        sections = [{"section": "FULL_TEXT", "text": text.strip()}]

    return sections


# ---------------------------
# 5) DOCX EXTRACTION (paragraphs + tables)
# ---------------------------
def extract_docx_text(docx_path: Path) -> str:
    doc = Document(docx_path)
    parts = []

    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text and p.text.strip():
                        parts.append(p.text.strip())

    return "\n".join(parts)


# ---------------------------
# 6) PROCESSING
# ---------------------------
def process_folder(input_dir: Path, output_dir: Path, recursive: bool = True) -> int:
    output_dir.mkdir(parents=True, exist_ok=True)
    pattern = "**/*.docx" if recursive else "*.docx"
    files = list(input_dir.glob(pattern))

    count = 0
    for docx_file in files:
        raw = extract_docx_text(docx_file)
        cleaned = clean_text(raw)
        sections = split_into_sections(cleaned)

        result = {
            "file": docx_file.name,
            "relative_path": str(docx_file.relative_to(input_dir)),
            "sections": sections,
        }

        out_path = output_dir / f"{docx_file.stem}.json"
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)

        count += 1

    return count


def main():
    parser = argparse.ArgumentParser(description="Clean and segment DOCX contracts into JSON.")
    parser.add_argument("--input", required=True, help="Input folder with .docx files")
    parser.add_argument("--output", required=True, help="Output folder for .json files")
    parser.add_argument("--no-recursive", action="store_true", help="Do not scan subfolders")
    args = parser.parse_args()

    input_dir = Path(args.input)
    output_dir = Path(args.output)

    if not input_dir.exists() or not input_dir.is_dir():
        raise SystemExit(f"Input folder not found: {input_dir}")

    processed = process_folder(input_dir, output_dir, recursive=not args.no_recursive)
    print(f"Done ✅ Processed {processed} file(s). Output: {output_dir}")


if __name__ == "__main__":
    main()
