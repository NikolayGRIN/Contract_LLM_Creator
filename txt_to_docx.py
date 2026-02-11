from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


TXT_PATH = Path("out.txt")
DOCX_PATH = Path("contract.docx")


# -----------------------------
# Styles / font normalization
# -----------------------------
def setup_docx_styles(doc: Document, font_name: str = "Times New Roman", font_size_pt: int = 12) -> None:
    """
    Force consistent fonts so Word doesn't switch fonts between paragraphs/runs
    (common when text mixes кириллица/latin or contains odd punctuation).
    """
    # Normal
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(font_size_pt)
    normal.element.rPr.rFonts.set(qn("w:ascii"), font_name)
    normal.element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    normal.element.rPr.rFonts.set(qn("w:cs"), font_name)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    # Headings (make them consistent too)
    for style_name, size, bold in [("Heading 1", 14, True), ("Heading 2", 12, True)]:
        if style_name in doc.styles:
            st = doc.styles[style_name]
            st.font.name = font_name
            st.font.size = Pt(size)
            st.font.bold = bold
            st.element.rPr.rFonts.set(qn("w:ascii"), font_name)
            st.element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
            st.element.rPr.rFonts.set(qn("w:cs"), font_name)
            st.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_paragraph_with_style(doc: Document, text: str, style_name: str) -> None:
    p = doc.add_paragraph(text)
    if style_name in doc.styles:
        p.style = doc.styles[style_name]
    else:
        p.style = doc.styles["Normal"]


# -----------------------------
# Heuristics: headings vs body
# -----------------------------
RE_SECTION_HEADING = re.compile(r"^\s*(\d+)\.\s+.+$")            # "1. ОПРЕДЕЛЕНИЯ"
RE_SUBCLAUSE = re.compile(r"^\s*(\d+(\.\d+)+)\.\s+.+$")         # "1.1. ...", "10.12. ..."
RE_ARTICLE = re.compile(r"^\s*(ARTICLE|SECTION)\b", re.I)       # "ARTICLE 10", "SECTION 5"


def classify_line(line: str) -> str:
    """
    Returns: "blank" | "heading1" | "normal"
    """
    if not line.strip():
        return "blank"

    s = line.strip()

    # ARTICLE/SECTION headings
    if RE_ARTICLE.match(s):
        return "heading1"

    # Main numbered headings: "1. TITLE"
    if RE_SECTION_HEADING.match(s) and not RE_SUBCLAUSE.match(s):
        return "heading1"

    return "normal"


def txt_to_docx(txt_path: Path, docx_path: Path) -> None:
    text = txt_path.read_text(encoding="utf-8")

    doc = Document()
    setup_docx_styles(doc, font_name="Times New Roman", font_size_pt=12)

    for raw in text.splitlines():
        # Preserve empty lines as paragraph breaks
        kind = classify_line(raw)

        if kind == "blank":
            doc.add_paragraph("")  # keeps spacing similar to TXT
            continue

        line = raw.strip()

        if kind == "heading1":
            add_paragraph_with_style(doc, line, "Heading 1")
        else:
            add_paragraph_with_style(doc, line, "Normal")

    doc.save(docx_path)


if __name__ == "__main__":
    txt_to_docx(TXT_PATH, DOCX_PATH)

